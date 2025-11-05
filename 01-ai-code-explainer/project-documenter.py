import os
import docx
from docx import Document
from docx.shared import Inches
import tkinter as tk
from tkinter import filedialog

def selecionar_pasta():
    """Seleciona a pasta raiz do projeto"""
    root = tk.Tk()
    root.withdraw()
    pasta = filedialog.askdirectory(title="Selecione a pasta raiz do projeto")
    return pasta

def criar_documentacao_word(pasta_raiz, arquivo_saida="documentacao_projeto.docx"):
    """Cria documentação estruturada no Word"""
    
    doc = Document()
    
    # Título principal
    titulo = doc.add_heading('DOCUMENTAÇÃO DO PROJETO', 0)
    titulo.alignment = 1
    
    # Informações da pasta raiz
    doc.add_heading(f'Pasta Raiz: {os.path.basename(pasta_raiz)}', level=1)
    doc.add_paragraph(f'Caminho completo: {pasta_raiz}')
    doc.add_paragraph()
    
    # Função recursiva para percorrer a estrutura de pastas
    def processar_pasta(caminho_pasta, nivel=1):
        try:
            itens = sorted(os.listdir(caminho_pasta))
            
            for item in itens:
                caminho_completo = os.path.join(caminho_pasta, item)
                
                if os.path.isdir(caminho_completo):
                    # É uma pasta
                    doc.add_heading(f'📁 Pasta: {item}', level=nivel)
                    doc.add_paragraph(f'Localização: {caminho_completo}')
                    doc.add_paragraph()
                    
                    # Processar subpastas recursivamente
                    processar_pasta(caminho_completo, nivel + 1)
                    
                else:
                    # É um arquivo
                    processar_arquivo(caminho_completo, nivel)
                    
        except PermissionError:
            doc.add_paragraph(f'⚠️ Acesso negado à pasta: {caminho_pasta}')
        except Exception as e:
            doc.add_paragraph(f'❌ Erro ao processar {caminho_pasta}: {str(e)}')
    
    def processar_arquivo(caminho_arquivo, nivel):
        """Processa um arquivo individual"""
        nome_arquivo = os.path.basename(caminho_arquivo)
        extensao = os.path.splitext(nome_arquivo)[1].lower()
        
        # Adicionar cabeçalho do arquivo
        doc.add_heading(f'📄 Arquivo: {nome_arquivo}', level=nivel)
        doc.add_paragraph(f'Localização: {caminho_arquivo}')
        doc.add_paragraph(f'Tamanho: {os.path.getsize(caminho_arquivo)} bytes')
        
        # Ler e adicionar conteúdo de arquivos de código
        if extensao in ['.py', '.js', '.java', '.cpp', '.c', '.html', '.css', '.php', '.rb', '.go', '.rs', '.ts']:
            adicionar_codigo_arquivo(caminho_arquivo, doc)
        else:
            doc.add_paragraph('Tipo de arquivo: Binário/Outro (conteúdo não exibido)')
        
        doc.add_paragraph()
    
    def adicionar_codigo_arquivo(caminho_arquivo, documento):
        """Adiciona o conteúdo de arquivos de código fonte"""
        try:
            with open(caminho_arquivo, 'r', encoding='utf-8') as arquivo:
                conteudo = arquivo.read()
            
            documento.add_heading('Conteúdo do Arquivo:', level=4)
            
            # Criar uma tabela para melhor formatação do código
            tabela = documento.add_table(rows=1, cols=1)
            tabela.style = 'Light Grid'
            celula = tabela.rows[0].cells[0]
            
            # Adicionar o código preservando a formatação
            paragrafo = celula.paragraphs[0]
            run = paragrafo.add_run(conteudo)
            run.font.name = 'Consolas'  # Fonte monoespaçada para código
            
        except UnicodeDecodeError:
            documento.add_paragraph('❌ Erro: Não foi possível ler o arquivo (codificação não suportada)')
        except Exception as e:
            documento.add_paragraph(f'❌ Erro ao ler arquivo: {str(e)}')
    
    # Iniciar processamento
    doc.add_heading('ESTRUTURA DO PROJETO', level=1)
    processar_pasta(pasta_raiz)
    
    # Resumo final
    doc.add_heading('RESUMO DO PROJETO', level=1)
    
    # Estatísticas
    total_arquivos = 0
    total_pastas = 0
    arquivos_codigo = 0
    
    for raiz, pastas, arquivos in os.walk(pasta_raiz):
        total_pastas += len(pastas)
        total_arquivos += len(arquivos)
        for arquivo in arquivos:
            if os.path.splitext(arquivo)[1].lower() in ['.py', '.js', '.java', '.cpp', '.c', '.html', '.css', '.php', '.rb', '.go', '.rs', '.ts']:
                arquivos_codigo += 1
    
    doc.add_paragraph(f'• Total de pastas: {total_pastas}')
    doc.add_paragraph(f'• Total de arquivos: {total_arquivos}')
    doc.add_paragraph(f'• Arquivos de código: {arquivos_codigo}')
    doc.add_paragraph(f'• Outros arquivos: {total_arquivos - arquivos_codigo}')
    
    # Salvar documento
    doc.save(arquivo_saida)
    return arquivo_saida

def main():
    """Função principal"""
    print("=== GERADOR DE DOCUMENTAÇÃO PARA WORD ===")
    print("Este programa criará uma documentação estruturada do seu projeto.")
    print()
    
    # Selecionar pasta
    pasta_raiz = selecionar_pasta()
    if not pasta_raiz:
        print("Nenhuma pasta selecionada. Programa encerrado.")
        return
    
    print(f"Pasta selecionada: {pasta_raiz}")
    
    # Nome do arquivo de saída
    nome_projeto = os.path.basename(pasta_raiz) or "projeto"
    arquivo_saida = f"documentacao_{nome_projeto}.docx"
    
    # Criar documentação
    print("Criando documentação...")
    try:
        arquivo_gerado = criar_documentacao_word(pasta_raiz, arquivo_saida)
        print(f"✅ Documentação criada com sucesso: {arquivo_gerado}")
        print()
        print("A documentação inclui:")
        print("• Estrutura completa de pastas")
        print("• Listagem de todos os arquivos")
        print("• Conteúdo dos arquivos de código fonte")
        print("• Estatísticas do projeto")
        
    except Exception as e:
        print(f"❌ Erro ao criar documentação: {str(e)}")

if __name__ == "__main__":
    main()